from flask import Flask, render_template, request, redirect, url_for, send_from_directory
from first_website_data import get_build_id, fetch_all_grants
from analysis import analyze_website
from application_generation import analyze_grant_feasibility
from docx import Document
from flask import send_file
from datetime import datetime
import os


app = Flask(__name__)
matched = []  # Global variable to store matched grants
GENERATED_FOLDER = "generated"
build_no = None
grants = None

def get_grants():
    global build_no, grants
    if grants is None:
        try:
            build_no = get_build_id()
            grants = fetch_all_grants(build_no)
        except Exception as e:
            print("Error fetching grants:", e)
            grants = []
    return grants

@app.route('/')
@app.route('/home')
def index():
    return render_template('index.html')



@app.route('/search', methods=['GET', 'POST'])
def search():
    matched = []
    business_industries = [
        "Technology", "Healthcare", "Education", "Agriculture", "Manufacturing",
        "Creative Industries", "Construction", "Finance", "Retail", "Hospitality",
        "Transportation", "Energy", "Environment", "Tourism"
    ]
        
  
    business_sizes = [
        "Startup", "Micro (1-9 employees)", "Small (10-49 employees)",
        "Medium (50-249 employees)", "Large (250+ employees)", "Charity / Non-profit"
    ]
 
    funding_purposes = [
        "Research and Development", "Innovation", "Sustainability",
        "Hiring and Training", "Equipment Purchase", "Business Expansion",
        "Export and International Trade", "Digital Transformation",
        "Community Projects", "Startup Support", "Other"
    ]   
    
    if request.method == 'POST':
        business_industry = request.form.get('business_industry')
        business_size = request.form.get('business_size')
        funding_purpose = request.form.get('funding_purpose')
        amount_needed = request.form.get('amount_needed')
        area = request.form.get('area')
        
        print("==== User Submitted Form ====")
        print("Industry:", business_industry)
        print("Business Size:", business_size)
        print("Funding Purpose:", funding_purpose)
        print("Amount Needed:", amount_needed)
        print("Region:", area)
        print("=============================")
        
        analysis_data = {
            'business_sector': business_industry,
            'company_size': business_size,
            'funding_purpose': funding_purpose,
            'preferred_amount': amount_needed,
            'location': area,
        }
       
        matched = match_grants(analysis_data)
        return render_template('results.html', match_grants=matched, analysis_data=analysis_data, from_search=True)
    
    return render_template(
        'search.html',
        business_industries=business_industries,
        business_sizes=business_sizes,
        funding_purposes=funding_purposes,
        matched=matched
    )


def match_grants(analysis_data):
    grants = get_grants()
    sector = analysis_data['business_sector'].lower()
    purpose = analysis_data['funding_purpose'].lower()
    amount = analysis_data['preferred_amount']
    location_usr = analysis_data['location'].lower()  # Convert user location to lowercase

    min_amount,max_amount = amount.split('-')
    min_amount = int(min_amount)
    max_amount = int(max_amount)

    matched_grants = []
    
    for grant in grants:
        name = grant.get('grantName','').lower()
        grant_location = [loc.lower() for loc in grant.get('grantLocation',[])]  # Convert all locations to lowercase
        desc = grant.get('grantShortDescription','').lower()
        min_reward = grant.get('grantMinimumAward',0)
        max_reward = grant.get('grantMaximumAward',0)
        url = grant.get('url','')

        sector_match = sector in name or sector in desc
        purpose_match = purpose in name or purpose in desc
        location_match = any(location_usr in loc for loc in grant_location)  # Compare lowercase locations
        amount_match = (
            min_reward <= max_amount and max_reward >= min_amount
        )
        match_score=sum([sector_match,purpose_match,location_match,amount_match])
        if match_score>=2:
            matched_grants.append(grant)
    return matched_grants


@app.route('/analyzer', methods=['GET', 'POST'])
def analyzer():
    business_industries = [
        "Technology", "Healthcare", "Education", "Agriculture", "Manufacturing",
        "Creative Industries", "Construction", "Finance", "Retail", "Hospitality",
        "Transportation", "Energy", "Environment", "Tourism"
    ]
    business_sizes = [
        "Startup", "Micro (1-9 employees)", "Small (10-49 employees)",
        "Medium (50-249 employees)", "Large (250+ employees)", "Charity / Non-profit"
    ]
    funding_purposes = [
        "Research and Development", "Innovation", "Sustainability",
        "Hiring and Training", "Equipment Purchase", "Business Expansion",
        "Export and International Trade", "Digital Transformation",
        "Community Projects", "Startup Support", "Other"
    ]
    error = None
    grants = get_grants()

    if request.method == 'POST':
        action = request.form.get('action')

        if action == 'analyze':
            web_url = request.form.get('website_url')
            analysis_result = {}  # ✅ Define early to avoid UnboundLocalError

            if web_url:
                analysis_text = analyze_website(web_url)

                if isinstance(analysis_text, dict):
                    analysis_result = {
                        "business_sector": analysis_text.get("Industry", ""),
                        "funding_purpose": analysis_text.get("Funding Purpose", ""),
                        "preferred_amount": analysis_text.get("Amount Needed", ""),
                        "location": analysis_text.get("Region", "")
                    }
                elif isinstance(analysis_text, str):
                    lines = analysis_text.strip().split("\n")
                    for line in lines:
                        if ':' in line:
                            key, value = line.split(":", 1)
                            analysis_result[key.strip().lower().replace(" ", "_")] = value.strip()
                else:
                    error = "Website analysis returned an unexpected format."
            else:
                error = "Please enter a website URL to analyze."

            return render_template(
                'analyzer.html',
                editable_result=analysis_result,
                business_industries=business_industries,
                business_sizes=business_sizes,
                funding_purposes=funding_purposes,
                show_editable_form=False,
                error=error
            )

        elif action == 'show_edit':
            editable_result = {k: v for k, v in request.form.items() if k != 'action'}
            return render_template(
                'analyzer.html',
                editable_result=editable_result,
                business_industries=business_industries,
                business_sizes=business_sizes,
                funding_purposes=funding_purposes,
                show_editable_form=True,
                error=error
            )

        elif action in ['match', 'edit']:
            analysis_data = {k: v for k, v in request.form.items() if k != 'action'}
            required_fields = ['business_sector', 'company_size', 'funding_purpose', 'preferred_amount', 'location']
            missing = [field for field in required_fields if not analysis_data.get(field)]

            if missing:
                error = "All fields are required to match grants."
                return render_template(
                    'analyzer.html',
                    editable_result=analysis_data,
                    business_industries=business_industries,
                    business_sizes=business_sizes,
                    funding_purposes=funding_purposes,
                    show_editable_form=False,
                    error=error
                )

            matched = match_grants(analysis_data)
            return render_template(
                'results.html',
                match_grants=matched,
                analysis_data=analysis_data,
                from_analyzer=True
            )

    # Default GET
    return render_template(
        'analyzer.html',
        business_industries=business_industries,
        business_sizes=business_sizes,
        funding_purposes=funding_purposes,
        show_editable_form=False,
        error=error
    )

def get_grant_by_id(grants,grant_id):
    for grant in grants:
        if grant.get('id') == grant_id:
            return grant
    return None
    

@app.route('/application/<grant_id>', methods=['GET', 'POST'])
def application(grant_id):
    grants = get_grants()
    grant_data = get_grant_by_id(grants, grant_id)
    if not grant_data:
        return "Grant not found"

    filename = None
    if request.method == 'POST':
        form_data = {
            'business_name': request.form.get('business_name'),
            'business_registration_number': request.form.get('business_registration_number'),
            'business_website': request.form.get('business_website'),
            'business_industry': request.form.get('business_industry'),
            'contact_name': request.form.get('contact_name'),
            'contact_email': request.form.get('contact_email'),
            'contact_phone': request.form.get('contact_phone'),
            'contact_address': request.form.get('contact_address'),
            'contact_state': request.form.get('contact_state'),
            'project_title': request.form.get('project_title'),
            'project_description': request.form.get('project_description'),
            'start_date': request.form.get('start_date'),
            'duration': request.form.get('duration'),
            'funding_amount': request.form.get('funding_amount'),
        }
        combined_data = grant_data.copy()
        combined_data.update(form_data)
        print(combined_data)
        gpt_analysis = analyze_grant_feasibility(combined_data)

        docx = Document()
        docx.add_heading("Grant Application")
        docx.add_paragraph(gpt_analysis)

        filename = f"grant_application_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        os.makedirs(GENERATED_FOLDER, exist_ok=True)
        filepath = os.path.join(GENERATED_FOLDER, filename)
        docx.save(filepath)

        return redirect(url_for('application', grant_id=grant_id, filename=filename))

    filename = request.args.get('filename')
    return render_template('application.html', grant=grant_data, filename=filename)

@app.route('/download/<filename>')
def download_file(filename):
    return send_from_directory(GENERATED_FOLDER, filename, as_attachment=True)


if __name__ == '__main__':
    app.run(debug=True)